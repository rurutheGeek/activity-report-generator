import random
import datetime
import docx
from docx.shared import Pt
from docx.enum.text import WD_UNDERLINE
import os
import re
import subprocess
import platform
import time
import sys
import shutil


def is_running_as_executable():
    """
    PyInstallerで作成された実行ファイルとして実行されているかどうかをチェック
    """
    return getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS')

def get_resource_path(relative_path):
    """
    リソースファイルの絶対パスを取得する関数
    PyInstaller実行の場合は_MEIPASSからのパス、通常実行の場合は相対パス
    """
    if is_running_as_executable():
        # PyInstallerで実行されている場合、_MEIPASSからのパスを返す
        base_path = sys._MEIPASS
    else:
        # 通常実行の場合、現在のディレクトリを基準にする
        base_path = os.path.abspath(".")
    
    return os.path.join(base_path, relative_path)

def main():
    # 現在の日付を取得
    current_date = datetime.datetime.now()
    current_year = current_date.year
    current_month = current_date.month
    
    # 活動月の入力
    print(f"活動報告更新プログラム\n")

    #活動年の入力
    while True:
        year_input = input(f"活動年を入力（{current_year}年で続行する場合はEnter）: ").strip()
        if not year_input:  # 入力が空の場合、現在の年を使用
            year = current_year
            break
        try:
            year = int(year_input)
        except ValueError:
            print("エラー: 有効な年を入力してください。")
    
    
    while True:
        month_input = input(f"活動月を入力（{current_month}月で続行する場合はEnter）: ").strip()
        if not month_input:  # 入力が空の場合、現在の月を使用
            month = current_month
            break
        try:
            month = int(month_input)
            if 1 <= month <= 12:
                break
            else:
                print("エラー: 1から12の間で入力してください。")
        except ValueError:
            print("エラー: 有効な月を入力してください。")
    
    # 場所の選択（八王子または新宿）
    while True:
        location_choice = input("活動場所を入力してください（h: 八王子, s: 新宿）: ").lower()
        if location_choice in ['h', 's']:
            break
        else:
            print("エラー: 'h'（八王子）または's'（新宿）を入力してください。")
            continue

    location_text = "（新宿）" if location_choice == "s" else "（八王子）"

    print(f"\n活動報告{location_text} - {year}年{month}月\n")
    
    # 活動日と場所の入力
    print("活動日と場所を最大6回入力してください")
    
    dates = []
    locations = []
    
    for i in range(6):
        date = input(f"活動日 {i+1} (例: 15、終了するにはEnter): ")
        if not date.strip():  # 日付が空白の場合、入力を終了
            break
            
        try:
            day = int(date)
            if 1 <= day <= 31:
                date = f"{month}/{day}"  # 月/日の形式に変換
            else:
                print("エラー: 1から31の間の日を入力してください。")
                continue
        except ValueError:
            print("エラー: 有効な日を入力してください。")
            continue
            
        location = input(f"活動場所 {i+1} (例: 02-264): ")
        dates.append(date)
        locations.append(location)
    
    # アクティビティのオプションと重みを設定
    activities = [
        ("ポケモンSV通信対戦", 30),
        ("ポケモンカードゲーム", 45),
        ("ポケカ開封会", 10),
        ("絵しりとり", 10),
        ("ホワイトボードアート", 10),
        ("情報交換", 10),
        ("ポケモンユナイト", 10),
        ("ポケモンソードシールド", 10),
        ("ポケモンポンジャン", 10),
        ("サークル会議", 0.3)
    ]
    
    #最小参加人数を設定
    minmember = 3
    #最大参加人数を設定
    maxmember = 14
    
    # 重みの正規化
    total_weight = sum(weight for _, weight in activities)
    normalized_activities = [(act, weight/total_weight) for act, weight in activities]
    
    # ランダムデータの生成
    report_data = []
    
    for i in range(len(dates)):
        # 参加者数をランダムに設定
        participants = random.randint(minmember, maxmember)
        
        # 重みに基づいて3つのアクティビティをランダムに選択
        selected_activities = random.choices(
            [act for act, _ in normalized_activities],
            weights=[weight for _, weight in normalized_activities],
            k=3
        )
        
        # 重複がないことを確認
        if len(set(selected_activities)) < 3:
            # 重複がある場合は別のアプローチを試す
            unique_activities = []
            available_activities = [act for act, _ in normalized_activities]
            
            while len(unique_activities) < 3 and available_activities:
                chosen = random.choices(
                    available_activities,
                    weights=[normalized_activities[available_activities.index(act)][1] 
                             for act in available_activities],
                    k=1
                )[0]
                unique_activities.append(chosen)
                available_activities.remove(chosen)
            
            selected_activities = unique_activities
        
        report_data.append({
            'date': dates[i],
            'location': locations[i],
            'participants': participants,
            'activities': selected_activities[:3]  # 3つのアクティビティのみを確保
        })
    
    # ドキュメントの読み込みと更新
    try:
        template_file = None
        
        if is_running_as_executable():
            # PyInstallerで実行されている場合、埋め込まれたリソースを使用
            template_file = get_resource_path("template/活動報告.docx")
            
            # 埋め込みリソースが存在するか確認
            if not os.path.exists(template_file):
                print(f"埋め込みテンプレートが見つかりません: {template_file}")
                
                # 代替方法として現在のディレクトリも検索
                files = os.listdir(".")
                for file in files:
                    if file.endswith("活動報告.docx"):
                        template_file = file
                        print(f"現在のディレクトリからテンプレートを読み込みます: {template_file}")
                        break
                        
                if not template_file:
                    print("テンプレートファイルが見つかりませんでした。")
                    return
        else:
            # 通常実行の場合、現在のディレクトリでテンプレートファイルを検索
            files = os.listdir(".")
            for file in files:
                if file.endswith("活動報告.docx"):
                    template_file = file
                    break
            
            if not template_file:
                print("テンプレートファイルが見つかりませんでした。")
                return
            
        print(f"テンプレートファイル: {template_file}")
        
        # ファイルが開かれているかチェック
        try:
            # ファイルを読み込む前にアクセスできるか確認
            with open(template_file, "rb") as f:
                pass
                
            doc = docx.Document(template_file)
        except PermissionError:
            print(f"エラー: {template_file}は現在開かれています。")
            print("ファイルを閉じてから再実行してください。")
            return
        except Exception as e:
            print(f"ファイルを開く際にエラーが発生しました: {e}")
            return
        
        # 年月の更新 - 特定のRunのみを更新
        updated = False
        
        if len(doc.paragraphs) > 0:
            # 最初の段落が年月情報を含む可能性が高い
            para = doc.paragraphs[0]
            
            # 年月の構造が分析結果と一致するかチェック
            if len(para.runs) >= 5 and '年' in para.text and '月' in para.text:
                
                # 指定されたRunのみを更新:
                if len(para.runs) >= 2:
                    year_str = str(year)
                    if len(year_str) >= 2:
                        para.runs[1].text = year_str[2:]  # 年の下2桁のみを更新
                
                # Run 4: '4' を更新 (月)
                if len(para.runs) >= 4:
                    para.runs[3].text = str(month)
                
                # Run 1, 3, 5 はテンプレートなので更新しない
                
                updated = True
        
        # 場所の更新（八王子/新宿）
        for para in doc.paragraphs:
            if 'サークル名　工学院ポケモンだいすきクラブ' in para.text:
                for run in para.runs:
                    if '（' in run.text and '）' in run.text:
                        run.text = location_text
                        break
        
        # アクティビティテーブルの更新
        if len(doc.tables) > 0:
            table = doc.tables[0]
            
            # 例の行を保持し、2行目から活動データを追加
            if len(table.rows) > 2:
                # 2行目（インデックス2、ヘッダーと例の後）から活動データを追加
                for i, data in enumerate(report_data):
                    # 2行目から開始（インデックス2、ヘッダーと例の後）
                    if i + 2 < len(table.rows):
                        row = table.rows[i + 2]
                    else:
                        # 6つ未満の活動では、ここに到達しないはず
                        break
                    
                    # 日付
                    if len(row.cells) > 0:
                        row.cells[0].text = data['date']
                    
                    # 場所
                    if len(row.cells) > 1:
                        row.cells[1].text = data['location']
                    
                    # 参加者
                    if len(row.cells) > 2:
                        row.cells[2].text = f"{data['participants']}人"
                    
                    # アクティビティ
                    if len(row.cells) > 3:
                        row.cells[3].text = "\n".join(data['activities'])
                
                # 残りの行をクリア（6つ未満の活動が入力された場合）
                for i in range(len(report_data) + 2, min(8, len(table.rows))):
                    for cell in table.rows[i].cells:
                        cell.text = ""
        
        if not updated:
            print("警告: 年月の更新が行われませんでした。")
        
        # ドキュメントを更新された年月とファイル名で保存
        output_dir = "./"
        new_filename = f"{output_dir}{year}年{month}月活動報告{location_text}.docx"
        
        # 絶対パスに変換
        abs_new_filename = os.path.abspath(new_filename)
        doc.save(new_filename)
        print(f"\n活動報告書が更新されました: {abs_new_filename}")
        
        # 絶対パスでファイルを開く試行
        try_open_file(abs_new_filename)
        
    except Exception as e:
        print(f"エラーが発生しました: {e}")
        import traceback
        traceback.print_exc()

def try_open_file(filepath):
    """
    生成されたファイルを既定のアプリケーションで開きます
    複数の方法を試みます
    """
    print(f"ファイルを開こうとしています: {filepath}")
    
    try:
        # 方法1: OSに応じたコマンドでファイルを開く
        if platform.system() == 'Darwin':  # macOS
            subprocess.call(('open', filepath))
            return True
        elif platform.system() == 'Windows':  # Windows
            # 方法1: os.startfile
            try:
                os.startfile(filepath)
                return True
            except Exception as e1:
                print(f"os.startfileでのオープンに失敗: {e1}")
                
                # 方法2: subprocessでコマンドを実行
                try:
                    subprocess.Popen(['start', filepath], shell=True)
                    return True
                except Exception as e2:
                    print(f"subprocessでのオープンに失敗: {e2}")
                    
                    # 方法3: 既定のアプリケーションを使用
                    try:
                        subprocess.call(['cmd', '/c', 'start', '', filepath])
                        return True
                    except Exception as e3:
                        print(f"cmd /c startでのオープンに失敗: {e3}")
        else:  # Linux
            subprocess.call(('xdg-open', filepath))
            return True
            
    except Exception as e:
        print(f"ファイルを開く際にエラーが発生しました: {e}")
    
    print("自動的にファイルを開くことができませんでした。")
    print(f"生成されたファイル '{filepath}' を手動で開いてください。")
    return False

if __name__ == "__main__":
    main()